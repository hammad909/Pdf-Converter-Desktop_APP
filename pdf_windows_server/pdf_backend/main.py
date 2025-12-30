from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pathlib import Path
import shutil
import uuid
import os
import pdfplumber
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from io import BytesIO
from PIL import Image
from txt_converter import router as txt_converter_router
from rtf_converter import router as rtf_converter_router
from pdf_converter import router as pdf_converter_router
from html_converter import router as html_converter_router
from convert_to_txt import router as txt_convert_to_router
from convert_to_rtf import router as rtf_convert_to_router
from convert_to_docx import router as docx_convert_to_router
from covert_to_ppt import router as ppt_convert_to_router
from convert_to_html import router as html_convert_to_router
from paths import APP_NAME, UPLOAD_DIR, CONVERTED_DIR
from contextlib import asynccontextmanager
import shutil
import socket
import tempfile
import uvicorn

@asynccontextmanager
async def lifespan(app: FastAPI):
    print("Server lifespan started")
    yield
    remove_port_file()


app = FastAPI(lifespan=lifespan)

app.include_router(txt_converter_router, prefix="/pdf")
app.include_router(rtf_converter_router, prefix="/pdf")
app.include_router(pdf_converter_router, prefix="/pdf")
app.include_router(html_converter_router, prefix="/pdf")

app.include_router(txt_convert_to_router, prefix="/txt")
app.include_router(rtf_convert_to_router, prefix="/rtf")
app.include_router(docx_convert_to_router, prefix="/docx")
app.include_router(ppt_convert_to_router, prefix="/pptx")
app.include_router(html_convert_to_router, prefix="/html")

@app.get("/")
def read_root():
    return {"message": "Welcome to PDF Converter"}

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], 
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def group_words_into_lines(words, y_tolerance=3):
    lines = []
    current_line = []
    current_y = None

    for word in sorted(words, key=lambda w: (w['top'], w['x0'])):
        if current_y is None or abs(word['top'] - current_y) <= y_tolerance:
            current_line.append(word)
            current_y = word['top']
        else:
            lines.append(current_line)
            current_line = [word]
            current_y = word['top']
    if current_line:
        lines.append(current_line)
    return lines

def extract_pdf_to_docx(pdf_path: str, output_path: str):
    doc = Document()

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            images = page.images
            image_entries = []

            for img in images:
                try:
                    x0, top, x1, bottom = img["x0"], img["top"], img["x1"], img["bottom"]
                    bbox = (x0, top, x1, bottom)
                    cropped_image = page.crop(bbox).to_image(resolution=150).original

                    img_byte_arr = BytesIO()
                    cropped_image.save(img_byte_arr, format="PNG")
                    img_byte_arr.seek(0)

                    image_entries.append({
                        "type": "image",
                        "top": top,
                        "data": img_byte_arr
                    })
                except Exception as e:
                    print(f"Image extraction failed: {e}")


            words = page.extract_words(extra_attrs=["fontname", "size"])
            text_lines = []
            if words:
                grouped_lines = group_words_into_lines(words)
                for line in grouped_lines:
                    text_lines.append({
                        "type": "text",
                        "top": line[0]['top'],
                        "data": line
                    })

            content = sorted(text_lines + image_entries, key=lambda item: item["top"])

            for item in content:
                if item["type"] == "text":
                    para = doc.add_paragraph()
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for word in item["data"]:
                        run = para.add_run(word["text"] + " ")
                    
                        if "size" in word:
                            run.font.size = Pt(float(word["size"]))
                    
                        fontname = word.get("fontname", "").lower()
                        if "bold" in fontname:
                            run.bold = True
                        if "italic" in fontname or "oblique" in fontname:
                            run.italic = True
                elif item["type"] == "image":
                    
                    try:
                        doc.add_picture(item["data"], width=Inches(4.5))
                        doc.add_paragraph()  
                    except Exception as e:
                        print(f"Failed to insert image: {e}")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

@app.post("/upload/")
async def upload_and_convert_pdf(file: UploadFile = File(...)):
    if file.content_type != "application/pdf":
        raise HTTPException(status_code=400, detail="Only PDF files are allowed")

    original_filename = Path(file.filename).stem
    file_id = str(uuid.uuid4())
    pdf_path = UPLOAD_DIR / f"{file_id}.pdf"
    with pdf_path.open("wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

   
    docx_filename = f"{original_filename}.docx"
    docx_path = CONVERTED_DIR / docx_filename


    try:
        extract_pdf_to_docx(str(pdf_path), str(docx_path))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {e}")

    return {
        "message": "PDF converted successfully",
        "file": docx_filename,
        "download_url": f"/download/{docx_filename}"
    }

@app.get("/download/{file_name}")
def download_file(file_name: str):
    file_path = CONVERTED_DIR / file_name 

    if file_path.exists():
        return FileResponse(
            path=file_path,
            filename=file_name,
            media_type="application/octet-stream"
        )

    raise HTTPException(status_code=404, detail="File not found")


# picking a free port 
def get_free_port():
    s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    s.bind(('', 0)) 
    port = s.getsockname()[1]
    s.close()
    return port

PORT_FILE = Path(os.getenv("LOCALAPPDATA")) / APP_NAME / "pdf_converter_port.txt"
PORT_FILE.parent.mkdir(parents=True, exist_ok=True)

def write_port_file(port: int):
    with open(PORT_FILE, "w") as f:
        f.write(str(port))
    print(f"Port {port} written to {PORT_FILE}")

def remove_port_file():
    try:
        if PORT_FILE.exists():
            PORT_FILE.unlink()
            print("Port file removed on shutdown.")
    except Exception as e:
        print(f"Failed to remove port file: {e}")



if __name__ == "__main__":
    port = get_free_port()
    write_port_file(port)

    uvicorn.run(
        app,
        host="127.0.0.1",
        port=port,
        log_config=None,
        access_log=False,
    )